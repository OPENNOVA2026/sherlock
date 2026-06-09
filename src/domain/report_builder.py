from collections.abc import Sequence
from datetime import datetime
from itertools import count
from typing import Optional
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

from src.domain.dataclasses import (
    CoordinatedAuthor,
    CoordinationAnalysis,
    CoordinationGroup,
    PublicMetrics,
    ReportConfig,
    TextCatalog,
    TopMessage,
    UserProfile,
)


class DocxUtils:
    @staticmethod
    def ensure_hyperlink_style(doc: Document) -> None:
        styles = doc.styles
        if "Hyperlink" not in [s.name for s in styles]:
            style = styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
            style.font.color.rgb = RGBColor(0, 0, 255)
            style.font.underline = True

    @staticmethod
    def add_hyperlink(paragraph, url: str, text: str) -> None:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "Hyperlink")
        r_pr.append(r_style)
        new_run.append(r_pr)

        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    @staticmethod
    def set_cell_margins(cell, top=0.08, start=0.08, bottom=0.08, end=0.08) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_mar = OxmlElement("w:tcMar")
        for margin, val in (
            ("top", top),
            ("start", start),
            ("bottom", bottom),
            ("end", end),
        ):
            node = OxmlElement(f"w:{margin}")
        # Convert inches to twips: 1 inch = 1440 twips
        twips = int(Inches(val).emu / 12700)
        node.set(qn("w:w"), str(twips))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
        tc_pr.append(tc_mar)


class CoordinationReportRenderer:
    def __init__(
        self, *, cfg: Optional[ReportConfig] = None, text: Optional[TextCatalog] = None
    ):
        self.cfg = cfg or ReportConfig()
        self.text = text or TextCatalog()

    # ---------- Public API ----------
    def render(self, analysis: CoordinationAnalysis) -> Document:
        doc = Document()
        DocxUtils.ensure_hyperlink_style(doc)

        self._render_title(doc)
        self._render_meta(doc)

        section_number = count(1)
        self._render_intro(next(section_number), doc, analysis)
        self._render_most(next(section_number), doc, analysis.most_coordinated_authors)
        doc.add_page_break()
        self._render_groups(next(section_number), doc, analysis.coordination_groups)
        self._render_methodology(next(section_number), doc)
        return doc

    # ---------- Section methods ----------
    def _render_title(self, doc: Document) -> None:
        title = doc.add_heading(self.text.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_meta(self, doc: Document) -> None:
        madrid_now = datetime.now(ZoneInfo(self.cfg.tz))
        p = doc.add_paragraph()
        p.add_run(
            f"{self.text.generated_on_prefix}{madrid_now.strftime('%d-%m-%Y %H:%M:%S')}"
        ).italic = True

    def _render_intro(
        self, idx: int, doc: Document, analysis: CoordinationAnalysis
    ) -> None:
        doc.add_heading(f"{idx}. {self.text.intro_title}", level=1)

        doc.add_paragraph(self.text.intro_body)
        doc.add_heading(f"{idx}.1 {self.text.intro_research}", level=2)

        # Render research facts if present, else placeholder
        if any(
            [
                analysis.request_title,
                analysis.request_description,
                analysis.request_query,
                analysis.time_from,
                analysis.time_to,
                analysis.inauthenticity_score is not None,
            ]
        ):
            facts = []
            if analysis.request_title:
                facts.append(f"Título: {analysis.request_title}")
            if analysis.request_description:
                facts.append(f"Descripción: {analysis.request_description}")
            if analysis.request_query:
                facts.append(f"Query: {analysis.request_query}")
            if analysis.time_from or analysis.time_to:
                from_t = analysis.time_from
                to_t = analysis.time_to
                facts.append(f"Horizonte: {from_t or '?'} → {to_t or '?'}")
            if analysis.inauthenticity_score is not None:
                score = analysis.inauthenticity_score
                facts.append(f"Coeficiente de comportamiento coordinado: {score}")
            for line in facts:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph(self.text.research_placeholder)

    def _render_user_inline(
        self, p, user: UserProfile, *, short: bool = False, two_cols: bool = False
    ) -> None:
        username = user.username or "unknown"
        display = user.display_name or ""
        desc = user.description or ""
        pm = user.metrics or PublicMetrics()
        followers = pm.followers_count if pm.followers_count is not None else "N/A"
        tweets = pm.tweet_count if pm.tweet_count is not None else "N/A"
        likes = pm.like_count if pm.like_count is not None else "N/A"

        DocxUtils.add_hyperlink(p, f"https://x.com/{username}", f"@{username}")
        if display:
            p.add_run(" — ")
            name_run = p.add_run(display)
            name_run.bold = True
        if not short:
            if desc:
                p.add_run(f"\n{desc}")
            stats = f"\nFollowers: {followers} • Tweets: {tweets}"
            if not two_cols:
                stats += f" • Likes: {likes}"
            else:
                stats += f"\nLikes: {likes}"
            p.add_run(stats)

    def _render_most(
        self, idx: int, doc: Document, most: Sequence[CoordinatedAuthor]
    ) -> None:
        doc.add_heading(f"{idx}. {self.text.most_title} ", level=1)

        doc.add_paragraph(self.text.most_preamble)
        if not most:
            doc.add_paragraph(self.text.most_empty)
            return

        table = doc.add_table(rows=0, cols=2)
        table.style = self.cfg.table_style

        header = table.add_row().cells
        for i, text in enumerate(["Información del usuario", "Grado de coord."]):
            p = header[i].paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        table.columns[0].width = Inches(self.cfg.user_col_width_inches)
        table.columns[1].width = Inches(self.cfg.degree_col_width_inches)

        for entry in most:
            row = table.add_row().cells
            # User info
            p1 = row[0].paragraphs[0]
            self._render_user_inline(p1, entry.user)
            row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            DocxUtils.set_cell_margins(row[0], *(self._margins_tuple()))
            # Degree
            p2 = row[1].paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run(
                str(entry.coordination_degree)
                if entry.coordination_degree is not None
                else "N/A"
            )
            row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            DocxUtils.set_cell_margins(row[1], *(self._margins_tuple()))

    def _render_groups(
        self, idx: int, doc: Document, groups: Sequence[CoordinationGroup]
    ) -> None:
        doc.add_heading(f"{idx}. {self.text.groups_title}", level=1)

        if not groups:
            doc.add_paragraph(self.text.groups_empty)
            return

        for gi, g in enumerate(groups, start=1):
            doc.add_heading(f"{idx}.{gi} Grupo {gi} con tamaño {g.size}", level=2)
            # Users in the group
            doc.add_heading(f"{idx}.{gi}.1 {self.text.groups_users_title}", level=3)
            self._render_users_two_columns(doc, g.users)

            # Pushed users
            doc.add_heading(
                f"{idx}.{gi}.2 {self.text.groups_pushed_users_title}", level=3
            )
            self._render_users_two_columns(doc, g.pushed_users)

            # Messages being pushed / top messages
            doc.add_heading(f"{idx}.{gi}.3 {self.text.groups_messages_title}", level=3)
            if g.top_messages:
                self._render_messages_table(doc, g.top_messages)
            else:
                doc.add_paragraph(self.text.groups_no_messages)
            doc.add_page_break()

    def _render_users_two_columns(
        self, doc: Document, users: Sequence[UserProfile]
    ) -> None:
        if not users:
            doc.add_paragraph("None")
            return
        table = doc.add_table(rows=0, cols=2)
        table.style = self.cfg.table_style
        table.autofit = True
        table.columns[0].width = Inches(self.cfg.two_col_width_inches)
        table.columns[1].width = Inches(self.cfg.two_col_width_inches)

        for i in range(0, len(users), 2):
            row = table.add_row().cells
            # Left cell
            p1 = row[0].paragraphs[0]
            self._render_user_inline(p1, users[i], two_cols=True)
            row[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            DocxUtils.set_cell_margins(row[0], *(self._margins_tuple()))
            # Right cell
            if i + 1 < len(users):
                p2 = row[1].paragraphs[0]
                self._render_user_inline(p2, users[i + 1], two_cols=True)
                row[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
                DocxUtils.set_cell_margins(row[1], *(self._margins_tuple()))

    def _render_messages_table(
        self, doc: Document, top_messages: Sequence[TopMessage]
    ) -> None:
        table = doc.add_table(rows=0, cols=3)

        table.style = self.cfg.table_style
        table.autofit = False
        table.columns[0].width = Inches(0.7)  # Count
        table.columns[1].width = Inches(
            self.cfg.message_author_col_width_inches
        )  # Author
        table.columns[2].width = Inches(
            self.cfg.message_text_col_width_inches
        )  # Message

        header = table.add_row().cells
        for i, text in enumerate(["Nº", "Autor", "Mensaje"]):
            p = header[i].paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for tm in top_messages:
            row = table.add_row().cells
            # Count
            p0 = row[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p0.add_run(str(tm.count))
            row[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            DocxUtils.set_cell_margins(row[0], *(self._margins_tuple()))

            # Author (username + hyperlink)
            p1 = row[1].paragraphs[0]
            author = tm.message.author or UserProfile(username="unknown")
            self._render_user_inline(p1, author, short=True)
            row[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            DocxUtils.set_cell_margins(row[1], *(self._margins_tuple()))

            # Message text (with link if available)
            p2 = row[2].paragraphs[0]
            if tm.message.external_url:
                DocxUtils.add_hyperlink(p2, tm.message.external_url, "Enlace")
                p2.add_run(": ")
            p2.add_run(tm.message.text)
            row[2].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            DocxUtils.set_cell_margins(row[2], *(self._margins_tuple()))

    def _render_methodology(self, idx: int, doc: Document) -> None:
        doc.add_heading(f"{idx}. {self.text.methodology_title}", level=1)
        doc.add_paragraph(self.text.methodology_body)

    # ---------- Misc helpers ----------
    def _margins_tuple(self):
        m = self.cfg.default_cell_margin_inches
        return (m, m, m, m)
